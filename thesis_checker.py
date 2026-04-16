from __future__ import annotations

import re
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Length, Pt
from lxml import etree
from pypdf import PdfReader

WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

ALLOWED_CHINESE_FONTS = {"DFKai-SB", "\u6a19\u6977\u9ad4"}
ALLOWED_ENGLISH_FONTS = {"Times New Roman"}

CHAPTER_PATTERN = re.compile(r"^\u7b2c[\u4e00-\u9fff0-9]+\u7ae0")
SECTION_PATTERN = re.compile(r"^(?:\u7b2c[\u4e00-\u9fff0-9]+\u7bc0|[0-9]+-[0-9]+(?:-[0-9]+)?)$")
ABSTRACT_HEADINGS = {"ABSTRACT", "\u6458\u8981"}
FRONT_HEADINGS = {
    "\u81f4\u8b1d",
    "\u8b1d\u8a8c",
    "\u76ee\u9304",
    "\u8868\u76ee\u9304",
    "\u5716\u76ee\u9304",
    "\u53c3\u8003\u6587\u737b",
}
KEYWORD_PREFIXES = ("Keywords", "\u95dc\u9375\u5b57")


@dataclass
class Issue:
    severity: str
    category: str
    title: str
    details: str
    location: str | None = None
    suggestion: str | None = None


def xml_bool(element, attr_name: str = "val") -> bool | None:
    if element is None:
        return None
    value = element.get(f"{{{WORD_NS['w']}}}{attr_name}")
    if value is None:
        return True
    return value not in {"0", "false", "False", "off"}


def style_id_from_paragraph(paragraph) -> str | None:
    ppr = getattr(paragraph._p, "pPr", None)
    pstyle = getattr(ppr, "pStyle", None)
    if pstyle is None:
        return None
    return pstyle.val


def style_id_from_run(run) -> str | None:
    rpr = getattr(run._r, "rPr", None)
    rstyle = getattr(rpr, "rStyle", None)
    if rstyle is None:
        return None
    return rstyle.val


def styles_root(document: Document):
    try:
        return document.part.styles_part.element
    except Exception:
        return None


def find_style_element(document: Document, style_id: str | None):
    if not style_id:
        return None
    root = styles_root(document)
    if root is None:
        return None
    xpath = f".//w:style[@w:styleId='{style_id}']"
    matches = root.xpath(xpath, namespaces=WORD_NS)
    return matches[0] if matches else None


def style_based_on_id(style_element) -> str | None:
    if style_element is None:
        return None
    based_on = style_element.find("w:basedOn", namespaces=WORD_NS)
    if based_on is None:
        return None
    return based_on.get(f"{{{WORD_NS['w']}}}val")


def iter_style_elements(document: Document, style_id: str | None):
    current_id = style_id
    visited = set()
    while current_id and current_id not in visited:
        visited.add(current_id)
        style_element = find_style_element(document, current_id)
        if style_element is None:
            break
        yield style_element
        current_id = style_based_on_id(style_element)


def xml_alignment_to_name(value: str | None) -> str | None:
    mapping = {
        "left": "\u9760\u5de6\u5c0d\u9f4a",
        "center": "\u7f6e\u4e2d",
        "right": "\u9760\u53f3\u5c0d\u9f4a",
        "both": "\u5de6\u53f3\u5c0d\u9f4a",
        "distribute": "\u5206\u6563\u5c0d\u9f4a",
    }
    return mapping.get(value)


def length_to_cm(length: Length | None) -> float | None:
    return None if length is None else round(length.cm, 2)


def pt_value(value) -> float | None:
    if value is None:
        return None
    try:
        return round(value.pt, 1)
    except Exception:
        return None


def alignment_name(value: int | None) -> str:
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "\u9760\u5de6\u5c0d\u9f4a",
        WD_ALIGN_PARAGRAPH.CENTER: "\u7f6e\u4e2d",
        WD_ALIGN_PARAGRAPH.RIGHT: "\u9760\u53f3\u5c0d\u9f4a",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "\u5de6\u53f3\u5c0d\u9f4a",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "\u5206\u6563\u5c0d\u9f4a",
    }
    return mapping.get(value, "\u672a\u6307\u5b9a")


def paragraph_text(paragraph) -> str:
    return re.sub(r"\s+", " ", paragraph.text or "").strip()


def visible_paragraphs(document: Document):
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph_text(paragraph)
        if text:
            yield index, paragraph, text


def run_font_name(run) -> str | None:
    if run.font.name:
        return run.font.name
    rfonts = getattr(getattr(run._element, "rPr", None), "rFonts", None)
    if rfonts is None:
        return None
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        value = rfonts.get(f"{{{WORD_NS['w']}}}{attr}")
        if value:
            return value
    return None


def iter_runs_with_text(paragraph) -> Iterable:
    for run in paragraph.runs:
        if run.text and run.text.strip():
            yield run


def iter_style_chain(style) -> Iterable:
    current = style
    while current is not None:
        yield current
        current = getattr(current, "base_style", None)


def direct_paragraph_alignment_name(paragraph) -> str | None:
    ppr = getattr(paragraph._p, "pPr", None)
    jc = getattr(ppr, "jc", None)
    if jc is not None and getattr(jc, "val", None):
        return xml_alignment_to_name(jc.val)
    return None


def effective_paragraph_alignment_name(document: Document, paragraph) -> str | None:
    direct_name = direct_paragraph_alignment_name(paragraph)
    if direct_name:
        return direct_name

    style = getattr(paragraph, "style", None)
    for current_style in iter_style_chain(style):
        alignment = current_style.paragraph_format.alignment
        if alignment is not None:
            return alignment_name(alignment)

    for style_element in iter_style_elements(document, style_id_from_paragraph(paragraph)):
        ppr = style_element.find("w:pPr", namespaces=WORD_NS)
        jc = ppr.find("w:jc", namespaces=WORD_NS) if ppr is not None else None
        if jc is not None:
            return xml_alignment_to_name(jc.get(f"{{{WORD_NS['w']}}}val"))

    text = paragraph_text(paragraph)
    raw_text = paragraph.text or ""
    if text and raw_text.count("\t") >= 1:
        return "\u7591\u4f3c\u4ee5 Tab \u505a\u8996\u89ba\u7f6e\u4e2d"
    return None


def effective_run_bold(run) -> bool | None:
    rpr = getattr(run._r, "rPr", None)
    if rpr is not None and getattr(rpr, "b", None) is not None:
        return xml_bool(rpr.b)
    if run.bold is not None:
        return run.bold

    style = getattr(run, "style", None)
    for current_style in iter_style_chain(style):
        if current_style.font.bold is not None:
            return current_style.font.bold
    return None


def paragraph_style_bold(document: Document, paragraph) -> bool | None:
    style = getattr(paragraph, "style", None)
    for current_style in iter_style_chain(style):
        if current_style.font.bold is not None:
            return current_style.font.bold

    for style_element in iter_style_elements(document, style_id_from_paragraph(paragraph)):
        rpr = style_element.find("w:rPr", namespaces=WORD_NS)
        bold = rpr.find("w:b", namespaces=WORD_NS) if rpr is not None else None
        if bold is not None:
            return xml_bool(bold)
    return None


def run_style_bold(document: Document, run) -> bool | None:
    for style_element in iter_style_elements(document, style_id_from_run(run)):
        rpr = style_element.find("w:rPr", namespaces=WORD_NS)
        bold = rpr.find("w:b", namespaces=WORD_NS) if rpr is not None else None
        if bold is not None:
            return xml_bool(bold)
    return None


def effective_paragraph_bold(document: Document, paragraph) -> bool:
    runs = list(iter_runs_with_text(paragraph))
    if not runs:
        return False
    bold_votes = 0
    known_votes = 0
    for run in runs:
        bold = effective_run_bold(run)
        if bold is None:
            bold = run_style_bold(document, run)
        if bold is None:
            bold = paragraph_style_bold(document, paragraph)
        if bold is not None:
            known_votes += 1
            if bold:
                bold_votes += 1
    if known_votes:
        return bold_votes >= max(1, known_votes // 2)
    paragraph_level = paragraph_style_bold(document, paragraph)
    return bool(paragraph_level)


def run_style_size(document: Document, run) -> float | None:
    for style_element in iter_style_elements(document, style_id_from_run(run)):
        rpr = style_element.find("w:rPr", namespaces=WORD_NS)
        sz = rpr.find("w:sz", namespaces=WORD_NS) if rpr is not None else None
        if sz is not None:
            value = sz.get(f"{{{WORD_NS['w']}}}val")
            if value and value.isdigit():
                return round(int(value) / 2, 1)
    return None


def paragraph_style_size(document: Document, paragraph) -> float | None:
    style = getattr(paragraph, "style", None)
    for current_style in iter_style_chain(style):
        size = pt_value(current_style.font.size)
        if size is not None:
            return size
    for style_element in iter_style_elements(document, style_id_from_paragraph(paragraph)):
        rpr = style_element.find("w:rPr", namespaces=WORD_NS)
        sz = rpr.find("w:sz", namespaces=WORD_NS) if rpr is not None else None
        if sz is not None:
            value = sz.get(f"{{{WORD_NS['w']}}}val")
            if value and value.isdigit():
                return round(int(value) / 2, 1)
    return None


def paragraph_fonts(paragraph) -> set[str]:
    return {
        font
        for run in iter_runs_with_text(paragraph)
        for font in [run_font_name(run)]
        if font
    }


def paragraph_sizes(document: Document, paragraph) -> set[float]:
    sizes = set()
    for run in iter_runs_with_text(paragraph):
        direct_size = pt_value(run.font.size)
        if direct_size is not None:
            sizes.add(direct_size)
            continue
        style_size = run_style_size(document, run)
        if style_size is not None:
            sizes.add(style_size)
    if sizes:
        return sizes
    paragraph_level_size = paragraph_style_size(document, paragraph)
    if paragraph_level_size is not None:
        sizes.add(paragraph_level_size)
    return sizes


def contains_cjk(text: str) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff]", text))


def contains_ascii_letters_or_digits(text: str) -> bool:
    return bool(re.search(r"[A-Za-z0-9]", text))


def paragraph_line_spacing(paragraph) -> tuple[str, float | None]:
    fmt = paragraph.paragraph_format
    if fmt.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE:
        return "1.5 \u500d\u884c\u9ad8", 1.5
    if isinstance(fmt.line_spacing, Pt):
        return "\u56fa\u5b9a\u503c", round(fmt.line_spacing.pt, 1)
    if isinstance(fmt.line_spacing, float):
        return "\u591a\u500d\u884c\u8ddd", round(float(fmt.line_spacing), 2)
    if fmt.line_spacing is not None:
        try:
            return "\u591a\u500d\u884c\u8ddd", round(float(fmt.line_spacing), 2)
        except Exception:
            return "\u5df2\u8a2d\u5b9a", None
    return "\u672a\u6307\u5b9a", None


def extract_docx_xml(docx_path: Path, member: str) -> bytes | None:
    try:
        with zipfile.ZipFile(docx_path) as archive:
            return archive.read(member)
    except KeyError:
        return None


def list_docx_members(docx_path: Path, prefix: str) -> list[str]:
    with zipfile.ZipFile(docx_path) as archive:
        return [name for name in archive.namelist() if name.startswith(prefix)]


def parse_xml(data: bytes | None):
    return None if not data else etree.fromstring(data)


def element_local_name(element) -> str:
    return etree.QName(element).localname


def roman_number(value: int, upper: bool = False) -> str:
    pairs = [
        (1000, "m"),
        (900, "cm"),
        (500, "d"),
        (400, "cd"),
        (100, "c"),
        (90, "xc"),
        (50, "l"),
        (40, "xl"),
        (10, "x"),
        (9, "ix"),
        (5, "v"),
        (4, "iv"),
        (1, "i"),
    ]
    result = []
    current = value
    for number, symbol in pairs:
        while current >= number:
            result.append(symbol)
            current -= number
    text = "".join(result)
    return text.upper() if upper else text


def roman_to_int(text: str) -> int | None:
    values = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100, "D": 500, "M": 1000}
    current = text.upper().strip()
    if not current or any(ch not in values for ch in current):
        return None
    total = 0
    previous = 0
    for ch in reversed(current):
        value = values[ch]
        if value < previous:
            total -= value
        else:
            total += value
            previous = value
    return total


def format_page_label(page_number: int, fmt: str) -> str:
    if fmt == "upperRoman":
        return roman_number(page_number, upper=True)
    if fmt == "lowerRoman":
        return roman_number(page_number, upper=False)
    return str(page_number)


def parse_section_page_settings(sect_pr, previous_start: int, previous_fmt: str) -> tuple[int, str]:
    pg_num_type = sect_pr.find("w:pgNumType", namespaces=WORD_NS) if sect_pr is not None else None
    start = previous_start
    fmt = previous_fmt
    if pg_num_type is not None:
        start_attr = pg_num_type.get(f"{{{WORD_NS['w']}}}start")
        fmt_attr = pg_num_type.get(f"{{{WORD_NS['w']}}}fmt")
        if start_attr:
            start = int(start_attr)
        if fmt_attr:
            fmt = fmt_attr
    return start, fmt


def normalize_text_for_match(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def paragraph_page_details(docx_path: Path) -> tuple[dict[int, str], dict[int, str], dict[int, str]]:
    root = parse_xml(extract_docx_xml(docx_path, "word/document.xml"))
    if root is None:
        return {}, {}, {}

    body = root.find("w:body", namespaces=WORD_NS)
    if body is None:
        return {}, {}, {}

    physical_page_number = 1
    paragraph_index = 0
    physical_map: dict[int, int] = {}
    paragraph_text_map: dict[int, str] = {}
    page_text_map: dict[int, list[str]] = {}
    sections: list[dict] = []
    section_start_para = 1
    current_start = 1
    current_fmt = "decimal"

    for child in body:
        local_name = element_local_name(child)

        if local_name == "p":
            paragraph_index += 1
            physical_map[paragraph_index] = physical_page_number
            page_text_map.setdefault(physical_page_number, [])

            paragraph_text = "".join(child.xpath(".//w:t/text()", namespaces=WORD_NS)).strip()
            paragraph_text_map[paragraph_index] = paragraph_text
            if paragraph_text:
                page_text_map[physical_page_number].append(paragraph_text)

            rendered_breaks = child.xpath(".//w:lastRenderedPageBreak", namespaces=WORD_NS)
            manual_breaks = child.xpath(".//w:br[@w:type='page']", namespaces=WORD_NS)

            ppr = child.find("w:pPr", namespaces=WORD_NS)
            sect_pr = ppr.find("w:sectPr", namespaces=WORD_NS) if ppr is not None else None
            if sect_pr is not None:
                current_start, current_fmt = parse_section_page_settings(sect_pr, current_start, current_fmt)
                sections.append(
                    {
                        "start_para": section_start_para,
                        "end_para": paragraph_index,
                        "start_page_number": current_start,
                        "fmt": current_fmt,
                    }
                )
                section_start_para = paragraph_index + 1

            physical_page_number += len(rendered_breaks) + len(manual_breaks)

        elif local_name == "sectPr":
            current_start, current_fmt = parse_section_page_settings(child, current_start, current_fmt)
            sections.append(
                {
                    "start_para": section_start_para,
                    "end_para": paragraph_index,
                    "start_page_number": current_start,
                    "fmt": current_fmt,
                }
            )

    if not sections and paragraph_index:
        sections.append(
            {
                "start_para": 1,
                "end_para": paragraph_index,
                "start_page_number": 1,
                "fmt": "decimal",
            }
        )

    label_map: dict[int, str] = {}
    physical_label_map: dict[int, str] = {}
    for section in sections:
        if section["start_para"] > section["end_para"]:
            continue
        start_para = section["start_para"]
        start_physical_page = physical_map.get(start_para, 1)
        for para_idx in range(section["start_para"], section["end_para"] + 1):
            physical_page = physical_map.get(para_idx, start_physical_page)
            offset = physical_page - start_physical_page
            displayed_page = section["start_page_number"] + offset
            page_label = format_page_label(displayed_page, section["fmt"])
            label_map[para_idx] = page_label
            physical_label_map[physical_page] = page_label

    rendered_page_text_map = {
        physical_page: " ".join(lines)
        for physical_page, lines in page_text_map.items()
    }

    return label_map, paragraph_text_map, {
        physical_label_map.get(physical_page, str(physical_page)): rendered_page_text_map.get(physical_page, "")
        for physical_page in sorted(page_text_map)
    }


def page_number_info(docx_path: Path) -> dict[str, str | bool]:
    footer_files = list_docx_members(docx_path, "word/footer")
    if not footer_files:
        return {"present": False, "format": "\u7121"}
    found = False
    page_format = "\u963f\u62c9\u4f2f\u6578\u5b57"
    for member in footer_files:
        text = (extract_docx_xml(docx_path, member) or b"").decode("utf-8", errors="ignore")
        if "PAGE" in text:
            found = True
            if "ROMAN" in text:
                page_format = "\u7f85\u99ac\u6578\u5b57"
    return {"present": found, "format": page_format if found else "\u7121"}


def has_watermark(docx_path: Path) -> bool:
    for member in list_docx_members(docx_path, "word/header"):
        text = (extract_docx_xml(docx_path, member) or b"").decode("utf-8", errors="ignore")
        if any(marker in text for marker in ("PowerPlusWaterMarkObject", "w:pict", "v:shape", "w:txbxContent", "o:spid")):
            return True
    return False


def document_protection(docx_path: Path) -> dict[str, str | bool]:
    settings_xml = parse_xml(extract_docx_xml(docx_path, "word/settings.xml"))
    if settings_xml is None:
        return {"enabled": False, "mode": "\u672a\u77e5"}
    protection = settings_xml.find("w:documentProtection", namespaces=WORD_NS)
    if protection is None:
        return {"enabled": False, "mode": "\u7121"}
    edit_mode = protection.get(f"{{{WORD_NS['w']}}}edit", "\u672a\u77e5")
    enforcement = protection.get(f"{{{WORD_NS['w']}}}enforcement", "0") == "1"
    return {"enabled": enforcement, "mode": edit_mode}


def classify_paragraph(text: str) -> str:
    if text in ABSTRACT_HEADINGS:
        return "摘要標題"
    if text in FRONT_HEADINGS:
        return "前置標題"
    if CHAPTER_PATTERN.match(text):
        return "章標題"
    if SECTION_PATTERN.match(text):
        return "節標題"
    if any(text.startswith(prefix) for prefix in KEYWORD_PREFIXES):
        return "關鍵字"
    return "內文"


def is_catalog_chapter_entry(text: str) -> bool:
    normalized = re.sub(r"\s+", "", text)
    if not normalized:
        return False
    if not CHAPTER_PATTERN.match(normalized):
        return False
    has_leader = "." in text or "\t" in text or "…" in text or "..." in text or "．" in text
    has_page_tail = bool(re.search(r"(\d+|[ivxlcdmIVXLCDM]+)\s*$", text))
    return has_leader or has_page_tail


def is_catalog_entry(text: str) -> bool:
    if not text.strip():
        return False
    has_leader = "." in text or "\t" in text or "…" in text or "..." in text or "．" in text
    has_page_tail = bool(re.search(r"(\d+|[ivxlcdmIVXLCDM]+)\s*$", text))
    return has_leader and has_page_tail


def add_issue(
    issues: list[Issue],
    severity: str,
    category: str,
    title: str,
    details: str,
    location: str | None = None,
    suggestion: str | None = None,
) -> None:
    issues.append(Issue(severity, category, title, details, location, suggestion))


def analyze_sections(document: Document, issues: list[Issue]) -> list[dict]:
    results = []
    for idx, section in enumerate(document.sections, start=1):
        page_width = length_to_cm(section.page_width)
        page_height = length_to_cm(section.page_height)
        top = length_to_cm(section.top_margin)
        bottom = length_to_cm(section.bottom_margin)
        left = length_to_cm(section.left_margin)
        right = length_to_cm(section.right_margin)
        results.append(
            {
                "index": idx,
                "page_width_cm": page_width,
                "page_height_cm": page_height,
                "top_cm": top,
                "bottom_cm": bottom,
                "left_cm": left,
                "right_cm": right,
                "start_type": str(section.start_type),
            }
        )
        if (page_width, page_height) != (21.0, 29.7):
            add_issue(issues, "error", "\u7248\u9762\u8a2d\u5b9a", f"\u7b2c {idx} \u7bc0\u4e0d\u662f A4 \u7d19\u5f35", f"\u76ee\u524d\u5075\u6e2c\u70ba {page_width} x {page_height} \u516c\u5206\uff0c\u898f\u7bc4\u8981\u6c42 A4 21.0 x 29.7 \u516c\u5206\u3002", f"\u7b2c {idx} \u7bc0", "\u8acb\u5728 Word \u5c07\u9019\u4e00\u7bc0\u7684\u7d19\u5f35\u5927\u5c0f\u8abf\u6574\u70ba A4\u3002")
        for side, expected_value, current in (("\u4e0a", 2.5, top), ("\u4e0b", 2.5, bottom), ("\u5de6", 3.0, left), ("\u53f3", 2.0, right)):
            if current is None or abs(current - expected_value) > 0.11:
                add_issue(issues, "error", "\u7248\u9762\u8a2d\u5b9a", f"\u7b2c {idx} \u7bc0\u7684{side}\u908a\u754c\u4e0d\u7b26\u898f\u7bc4", f"\u76ee\u524d\u5075\u6e2c\u70ba {current} \u516c\u5206\uff0c\u898f\u7bc4\u8981\u6c42 {expected_value} \u516c\u5206\u3002", f"\u7b2c {idx} \u7bc0", "\u8acb\u8abf\u6574\u9019\u4e00\u7bc0\u7684\u908a\u754c\u8a2d\u5b9a\u3002")
        if idx > 1 and section.start_type != WD_SECTION_START.NEW_PAGE:
            add_issue(issues, "warning", "\u5206\u7bc0\u8d77\u59cb", f"\u7b2c {idx} \u7bc0\u672a\u660e\u78ba\u5f9e\u65b0\u9801\u958b\u59cb", "\u898f\u7bc4\u8981\u6c42\u4e3b\u8981\u7bc0\u8207\u7ae0\u7bc0\u61c9\u5f9e\u65b0\u9801\u958b\u59cb\u3002", f"\u7b2c {idx} \u7bc0", "\u8acb\u5c07\u8a72\u8655\u6539\u70ba\u300c\u4e0b\u4e00\u9801\u5206\u7bc0\u7b26\u865f\u300d\u3002")
    return results


def check_paragraphs(
    document: Document,
    issues: list[Issue],
    page_map: dict[int, str],
    paragraph_text_map: dict[int, str],
    page_text_map: dict[str, str],
) -> list[dict]:
    paragraph_summaries = []
    in_abstract = False
    in_catalog = False
    for index, paragraph, text in visible_paragraphs(document):
        kind = classify_paragraph(text)
        if kind == "摘要標題":
            in_abstract = True
        elif kind in {"章標題", "前置標題"} and text != "\u6458\u8981":
            in_abstract = False
        catalog_like = is_catalog_entry(text)
        catalog_chapter_like = is_catalog_chapter_entry(text)

        if text == "\u76ee\u9304":
            in_catalog = True
        elif kind == "章標題" and not catalog_chapter_like:
            in_catalog = False
        fonts = paragraph_fonts(paragraph)
        sizes = paragraph_sizes(document, paragraph)
        effective_alignment = effective_paragraph_alignment_name(document, paragraph)
        alignment = effective_alignment or "\u672a\u6307\u5b9a"
        line_spacing_label, line_spacing_value = paragraph_line_spacing(paragraph)
        page_number = page_map.get(index)
        source_text = paragraph_text_map.get(index, text)
        normalized_source = normalize_text_for_match(source_text)
        page_text = page_text_map.get(page_number or "", "")
        normalized_page_text = normalize_text_for_match(page_text)
        page_confirmed = bool(page_number and normalized_source and normalized_source in normalized_page_text)

        roman_value = roman_to_int(page_number) if page_number else None
        suspicious_roman = roman_value is not None and roman_value > 20

        if page_confirmed and not suspicious_roman:
            location = f"\u7b2c {page_number} \u9801"
        elif page_number:
            location = f"\u9801\u78bc\u5f85\u78ba\u8a8d\uff08\u76ee\u524d\u63a8\u5b9a\u70ba\u7b2c {page_number} \u9801\uff09"
        else:
            location = f"\u9801\u78bc\u7121\u6cd5\u78ba\u5b9a\uff08\u6bb5\u843d\u5e8f\u865f {index}\uff09"

        paragraph_summaries.append({"頁碼": page_number or "\u7121\u6cd5\u5224\u5b9a", "文字內容": text[:120], "段落類型": kind, "字型": sorted(fonts), "字級": sorted(sizes), "對齊": alignment, "行距": line_spacing_label})
        if catalog_chapter_like:
            if not effective_paragraph_bold(document, paragraph):
                add_issue(issues, "error", "\u76ee\u9304\u9801\u683c\u5f0f", f"{location}\u7684\u76ee\u9304\u7ae0\u9805\u76ee\u672a\u8a2d\u70ba\u7c97\u9ad4", f"\u5075\u6e2c\u5230\u7684\u76ee\u9304\u6587\u5b57\u70ba\u300c{text}\u300d\u3002", location, "\u76ee\u9304\u4e2d\u7684\u7ae0\u9805\u76ee\u8acb\u8a2d\u70ba 12 pt \u7c97\u9ad4\u3002")
            if sizes and 12.0 not in sizes:
                add_issue(issues, "error", "\u76ee\u9304\u9801\u683c\u5f0f", f"{location}\u7684\u76ee\u9304\u7ae0\u9805\u76ee\u4e0d\u662f 12 pt", f"\u5075\u6e2c\u5230\u7684\u5b57\u7d1a\u70ba {sorted(sizes)}\u3002", location, "\u76ee\u9304\u4e2d\u7684\u7ae0\u9805\u76ee\u8acb\u8a2d\u70ba 12 pt \u7c97\u9ad4\u3002")
            elif not sizes:
                add_issue(issues, "warning", "\u76ee\u9304\u9801\u683c\u5f0f", f"{location}\u7684\u76ee\u9304\u7ae0\u9805\u76ee\u5b57\u7d1a\u7121\u6cd5\u660e\u78ba\u5224\u5b9a", f"\u5075\u6e2c\u5230\u7684\u76ee\u9304\u6587\u5b57\u70ba\u300c{text}\u300d\u3002", location, "\u8acb\u78ba\u8a8d\u76ee\u9304\u4e2d\u7684\u7ae0\u9805\u76ee\u70ba 12 pt \u7c97\u9ad4\u3002")
            if alignment not in {"\u5de6\u53f3\u5c0d\u9f4a", "\u5206\u6563\u5c0d\u9f4a"}:
                add_issue(issues, "error", "\u76ee\u9304\u9801\u683c\u5f0f", f"{location}\u7684\u76ee\u9304\u7ae0\u9805\u76ee\u672a\u5de6\u53f3\u5c0d\u9f4a", f"\u5075\u6e2c\u5230\u7684\u5c0d\u9f4a\u65b9\u5f0f\u70ba\u300c{alignment}\u300d\u3002", location, "\u76ee\u9304\u4e2d\u7684\u7ae0\u9805\u76ee\u8acb\u8a2d\u70ba\u5de6\u53f3\u5c0d\u9f4a\u3002")
            continue
        if kind == "章標題":
            if alignment not in {"\u7f6e\u4e2d", "\u7591\u4f3c\u4ee5 Tab \u505a\u8996\u89ba\u7f6e\u4e2d"}:
                add_issue(issues, "error", "\u6a19\u984c\u683c\u5f0f", f"{location}\u7684\u7ae0\u6a19\u984c\u672a\u7f6e\u4e2d", f"\u5075\u6e2c\u5230\u7684\u6a19\u984c\u6587\u5b57\u70ba\u300c{text}\u300d\uff0c\u5c0d\u9f4a\u65b9\u5f0f\u70ba\u300c{alignment}\u300d\u3002", location, "\u8acb\u5c07\u7ae0\u6a19\u984c\u8a2d\u70ba\u7f6e\u4e2d\u3002")
            if not effective_paragraph_bold(document, paragraph):
                add_issue(issues, "error", "\u6a19\u984c\u683c\u5f0f", f"{location}\u7684\u7ae0\u6a19\u984c\u672a\u8a2d\u70ba\u7c97\u9ad4", f"\u5075\u6e2c\u5230\u7684\u6a19\u984c\u6587\u5b57\u70ba\u300c{text}\u300d\u3002", location, "\u8acb\u5c07\u7ae0\u6a19\u984c\u8a2d\u70ba\u7c97\u9ad4\u3002")
            if 16.0 not in sizes:
                add_issue(issues, "error", "\u6a19\u984c\u683c\u5f0f", f"{location}\u7684\u7ae0\u6a19\u984c\u4e0d\u662f 16 pt", f"\u5075\u6e2c\u5230\u7684\u5b57\u7d1a\u70ba {sorted(sizes) or '\u672a\u660e\u78ba\u8a2d\u5b9a'}\u3002", location, "\u8acb\u5c07\u7ae0\u6a19\u984c\u8a2d\u70ba 16 pt\u3002")
        elif kind == "節標題":
            if not effective_paragraph_bold(document, paragraph):
                add_issue(issues, "error", "\u6a19\u984c\u683c\u5f0f", f"{location}\u7684\u7bc0\u6a19\u984c\u672a\u8a2d\u70ba\u7c97\u9ad4", f"\u5075\u6e2c\u5230\u7684\u6a19\u984c\u6587\u5b57\u70ba\u300c{text}\u300d\u3002", location, "\u8acb\u5c07\u7bc0\u6a19\u984c\u8a2d\u70ba\u7c97\u9ad4\u3002")
            if 14.0 not in sizes:
                add_issue(issues, "error", "\u6a19\u984c\u683c\u5f0f", f"{location}\u7684\u7bc0\u6a19\u984c\u4e0d\u662f 14 pt", f"\u5075\u6e2c\u5230\u7684\u5b57\u7d1a\u70ba {sorted(sizes) or '\u672a\u660e\u78ba\u8a2d\u5b9a'}\u3002", location, "\u8acb\u5c07\u7bc0\u6a19\u984c\u8a2d\u70ba 14 pt\u3002")
        elif kind in {"摘要標題", "前置標題"}:
            if alignment not in {"\u7f6e\u4e2d", "\u7591\u4f3c\u4ee5 Tab \u505a\u8996\u89ba\u7f6e\u4e2d"}:
                add_issue(issues, "warning", "\u524d\u7f6e\u9801", f"{location}\u7684\u524d\u7f6e\u6a19\u984c\u672a\u7f6e\u4e2d", f"\u5075\u6e2c\u5230\u7684\u6a19\u984c\u6587\u5b57\u70ba\u300c{text}\u300d\uff0c\u5c0d\u9f4a\u65b9\u5f0f\u70ba\u300c{alignment}\u300d\u3002", location, "\u8acb\u5c07\u6a19\u984c\u8abf\u6574\u70ba\u7f6e\u4e2d\u3002")
        elif kind == "關鍵字":
            if 14.0 not in sizes:
                add_issue(issues, "warning", "\u6458\u8981\u8207\u95dc\u9375\u5b57", f"{location}\u7684\u95dc\u9375\u5b57\u6bb5\u843d\u4e0d\u662f 14 pt", f"\u5075\u6e2c\u5230\u7684\u5b57\u7d1a\u70ba {sorted(sizes) or '\u672a\u660e\u78ba\u8a2d\u5b9a'}\u3002", location, "\u8acb\u5c07\u95dc\u9375\u5b57\u6bb5\u843d\u8a2d\u70ba 14 pt\u3002")
        else:
            if in_abstract or kind == "內文":
                if contains_cjk(text) and fonts and not fonts.intersection(ALLOWED_CHINESE_FONTS):
                    add_issue(issues, "warning", "\u5167\u6587\u5b57\u578b", f"{location}\u7684\u4e2d\u6587\u5167\u6587\u5b57\u578b\u53ef\u80fd\u4e0d\u7b26\u898f\u7bc4", f"\u5075\u6e2c\u5230\u7684\u5b57\u578b\u70ba {', '.join(sorted(fonts))}\uff0c\u898f\u7bc4\u5efa\u8b70\u4e2d\u6587\u5167\u6587\u4f7f\u7528\u6a19\u6977\u9ad4 / DFKai-SB\u3002", location, "\u8acb\u5c07\u4e2d\u6587\u5167\u6587\u8abf\u6574\u70ba\u6a19\u6977\u9ad4 / DFKai-SB\u3002")
                if contains_ascii_letters_or_digits(text):
                    english_fonts = {run_font_name(run) for run in iter_runs_with_text(paragraph) if contains_ascii_letters_or_digits(run.text)}
                    english_fonts = {font for font in english_fonts if font}
                    if english_fonts and not english_fonts.issubset(ALLOWED_ENGLISH_FONTS):
                        add_issue(issues, "warning", "\u82f1\u6587\u8207\u6578\u5b57", f"{location}\u7684\u82f1\u6587\u6216\u6578\u5b57\u5b57\u578b\u53ef\u80fd\u4e0d\u7b26\u898f\u7bc4", f"\u5075\u6e2c\u5230\u7684\u5b57\u578b\u70ba {', '.join(sorted(english_fonts))}\uff0c\u898f\u7bc4\u5efa\u8b70\u82f1\u6587\u8207\u6578\u5b57\u4f7f\u7528 Times New Roman\u3002", location, "\u8acb\u5c07\u82f1\u6587\u8207\u6578\u5b57\u8abf\u6574\u70ba Times New Roman\u3002")
                if sizes and 12.0 not in sizes:
                    add_issue(issues, "warning", "\u5167\u6587\u5b57\u7d1a", f"{location}\u7684\u5167\u6587\u4e0d\u662f 12 pt", f"\u5075\u6e2c\u5230\u7684\u5b57\u7d1a\u70ba {sorted(sizes)}\uff0c\u898f\u7bc4\u8981\u6c42\u5167\u6587\u70ba 12 pt\u3002", location, "\u8acb\u5c07\u5167\u6587\u8abf\u6574\u70ba 12 pt\u3002")
                if line_spacing_value is None or abs(line_spacing_value - 1.5) > 0.05:
                    add_issue(issues, "warning", "\u884c\u8ddd", f"{location}\u672a\u4f7f\u7528 1.5 \u500d\u884c\u9ad8", f"\u5075\u6e2c\u5230\u7684\u884c\u8ddd\u70ba {line_spacing_label} {line_spacing_value or ''}".strip(), location, "\u8acb\u5c07\u6458\u8981\u8207\u5167\u6587\u6bb5\u843d\u8abf\u6574\u70ba 1.5 \u500d\u884c\u9ad8\u3002")
    return paragraph_summaries


def summarize_document(document: Document, docx_path: Path, issues: list[Issue]) -> dict:
    page_number = page_number_info(docx_path)
    watermark = has_watermark(docx_path)
    protection = document_protection(docx_path)
    if not page_number["present"]:
        add_issue(issues, "warning", "\u9801\u78bc", "\u672a\u5075\u6e2c\u5230\u9801\u78bc\u6b04\u4f4d", "\u898f\u7bc4\u8981\u6c42\u524d\u7f6e\u9801\u4f7f\u7528\u7f85\u99ac\u6578\u5b57\u9801\u78bc\uff0c\u6b63\u6587\u4f7f\u7528\u963f\u62c9\u4f2f\u6578\u5b57\u9801\u78bc\uff0c\u4f46\u76ee\u524d\u5728 footer XML \u4e2d\u672a\u627e\u5230 PAGE \u6b04\u4f4d\u3002", suggestion="\u8acb\u5728 Word \u9801\u5c3e\u63d2\u5165\u9801\u78bc\uff0c\u4e26\u5206\u5225\u8a2d\u5b9a\u524d\u7f6e\u9801\u8207\u6b63\u6587\u7684\u9801\u78bc\u683c\u5f0f\u3002")
    if not watermark:
        add_issue(issues, "warning", "\u6d6e\u6c34\u5370", "\u672a\u5075\u6e2c\u5230\u6d6e\u6c34\u5370\u7269\u4ef6", "\u898f\u7bc4\u8981\u6c42\u5f9e\u81f4\u8b1d\u9801\u958b\u59cb\u52a0\u5165\u6d6e\u6c34\u5370\uff0c\u4f46\u6b64\u5de5\u5177\u53ea\u80fd\u6aa2\u67e5 DOCX \u5167\u662f\u5426\u5b58\u5728\u985e\u4f3c\u6d6e\u6c34\u5370\u7684\u7269\u4ef6\u3002", suggestion="\u8acb\u5728 Word \u4e2d\u6aa2\u67e5\u5404\u7bc0 header\uff0c\u78ba\u8a8d\u6d6e\u6c34\u5370\u662f\u5426\u5f9e\u81f4\u8b1d\u9801\u958b\u59cb\u3002")
    if not protection["enabled"]:
        add_issue(issues, "info", "\u4fdd\u8b77\u8a2d\u5b9a", "\u672a\u5075\u6e2c\u5230 Word \u6587\u4ef6\u4fdd\u8b77", "\u5716\u66f8\u9928\u898f\u7bc4\u6709\u63d0\u5230\u6587\u4ef6\u4fdd\u5168\uff0c\u4f46\u9019\u9805\u8981\u6c42\u5e38\u5e38\u5957\u7528\u5728\u6700\u7d42\u4e0a\u50b3\u6a94\u6216 PDF\u3002\u76ee\u524d\u9019\u4efd DOCX \u5167\u672a\u627e\u5230\u5df2\u555f\u7528\u7684 documentProtection \u8a2d\u5b9a\u3002", suggestion="\u82e5\u5b78\u6821\u8981\u6c42 PDF \u4fdd\u8b77\uff0c\u8acb\u5728\u532f\u51fa\u6700\u7d42 PDF \u5f8c\u518d\u884c\u78ba\u8a8d\u3002")
    return {"paragraph_count": len(document.paragraphs), "section_count": len(document.sections), "page_number": page_number, "watermark": watermark, "protection": protection}


def analyze_docx(docx_path: str | Path) -> dict:
    path = Path(docx_path)
    document = Document(path)
    issues: list[Issue] = []
    page_map, paragraph_text_map, page_text_map = paragraph_page_details(path)
    section_results = analyze_sections(document, issues)
    paragraph_results = check_paragraphs(document, issues, page_map, paragraph_text_map, page_text_map)
    coverage = summarize_document(document, path, issues)
    severity_rank = {"error": 0, "warning": 1, "info": 2}
    issues_sorted = sorted(issues, key=lambda item: (severity_rank.get(item.severity, 9), item.category, item.location or ""))
    return {
        "file_type": "docx",
        "file_name": path.name,
        "summary": {
            "errors": sum(1 for item in issues_sorted if item.severity == "error"),
            "warnings": sum(1 for item in issues_sorted if item.severity == "warning"),
            "infos": sum(1 for item in issues_sorted if item.severity == "info"),
            "checked_items": ["A4 \u7d19\u5f35\u5927\u5c0f", "\u9801\u9762\u908a\u754c", "\u7ae0\u6a19\u984c\u683c\u5f0f", "\u7bc0\u6a19\u984c\u683c\u5f0f", "\u6458\u8981\u8207\u95dc\u9375\u5b57", "\u5167\u6587\u5b57\u578b\u8207\u5b57\u7d1a", "1.5 \u500d\u884c\u9ad8", "\u9801\u78bc\u6b04\u4f4d", "\u6d6e\u6c34\u5370\u5b58\u5728\u6027", "Word \u4fdd\u8b77\u8a2d\u5b9a"],
            "limitations": ["DOCX \u5206\u6790\u7121\u6cd5\u5b8c\u6574\u9084\u539f\u6700\u7d42\u5206\u9801\uff0c\u56e0\u6b64\u7121\u6cd5\u4fdd\u8b49\u6bcf\u7ae0\u90fd\u843d\u5728\u5947\u6578\u9801\u3002", "\u6d6e\u6c34\u5370\u6aa2\u67e5\u53ea\u80fd\u78ba\u8a8d DOCX \u5167\u662f\u5426\u5b58\u5728\u985e\u4f3c\u6d6e\u6c34\u5370\u7684\u7269\u4ef6\u3002", "\u82e5\u5b57\u578b\u50c5\u7531\u6a23\u5f0f\u6216\u4e3b\u984c\u7e7c\u627f\uff0c\u5de5\u5177\u53ef\u80fd\u9700\u8981\u505a\u8f03\u4fdd\u5b88\u7684\u5224\u65b7\u3002", "\u7d19\u5f35\u78c5\u6578\uff0c\u96d9\u9762\u5217\u5370\u8207\u6700\u7d42 PDF \u4fdd\u5168\u8a2d\u5b9a\u4ecd\u9700\u8981\u4eba\u5de5\u78ba\u8a8d\u3002"],
        },
        "coverage": coverage,
        "section_results": section_results,
        "paragraph_results": paragraph_results,
        "issues": [asdict(item) for item in issues_sorted],
    }


def detect_pdf_page_label(page_text: str) -> str | None:
    lines = [line.strip() for line in page_text.splitlines() if line.strip()]
    if not lines:
        return None
    candidates = lines[-3:]
    roman_pattern = re.compile(r"^[ivxlcdmIVXLCDM]+$")
    arabic_pattern = re.compile(r"^\d{1,4}$")
    for candidate in reversed(candidates):
        if arabic_pattern.match(candidate) or roman_pattern.match(candidate):
            return candidate
    return None


def analyze_pdf(pdf_path: str | Path) -> dict:
    path = Path(pdf_path)
    reader = PdfReader(str(path))
    issues: list[Issue] = []
    page_results = []
    detected_labels: list[str] = []
    total_lines = 0

    found_abstract = False
    found_catalog = False
    found_reference = False
    found_chapter = False

    for idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        total_lines += len(lines)
        page_label = detect_pdf_page_label(text)
        if page_label:
            detected_labels.append(page_label)

        joined = " ".join(lines[:40])
        if "\u6458\u8981" in joined or "ABSTRACT" in joined:
            found_abstract = True
        if "\u76ee\u9304" in joined:
            found_catalog = True
        if "\u53c3\u8003\u6587\u737b" in joined:
            found_reference = True
        if re.search(r"\u7b2c[\u4e00-\u9fff0-9]+\u7ae0", joined):
            found_chapter = True

        page_results.append(
            {
                "\u5be6\u9ad4\u9801": idx,
                "\u9801\u5c3e\u9801\u78bc\u7dda\u7d22": page_label or "\u672a\u5075\u6e2c\u5230",
                "\u53ef\u64f7\u53d6\u884c\u6578": len(lines),
                "\u5167\u5bb9\u9810\u89bd": joined[:120],
            }
        )

        if not lines:
            issues.append(
                Issue(
                    severity="warning",
                    category="\u9801\u9762\u5167\u5bb9",
                    title=f"\u5be6\u9ad4\u7b2c {idx} \u9801\u7121\u6cd5\u64f7\u53d6\u6587\u5b57",
                    details="\u9019\u4e00\u9801\u53ef\u80fd\u662f\u63c3\u63cf PDF\uff0c\u6216\u6587\u5b57\u5c64\u4e0d\u5b8c\u6574\uff0c\u56e0\u6b64\u7121\u6cd5\u9032\u884c\u7d30\u90e8\u683c\u5f0f\u6aa2\u67e5\u3002",
                    location=f"\u5be6\u9ad4\u7b2c {idx} \u9801",
                    suggestion="\u82e5\u8981\u505a\u66f4\u7cbe\u6e96\u7684\u683c\u5f0f\u6aa2\u67e5\uff0c\u8acb\u76e1\u91cf\u4e0a\u50b3 DOCX\uff0c\u6216\u78ba\u8a8d PDF \u5177\u6709\u53ef\u64f7\u53d6\u7684\u6587\u5b57\u5c64\u3002",
                )
            )

    if not detected_labels:
        issues.append(
            Issue(
                severity="warning",
                category="\u9801\u78bc",
                title="\u672a\u5f9e PDF \u9801\u5c3e\u5075\u6e2c\u5230\u660e\u986f\u9801\u78bc",
                details="\u76ee\u524d\u672a\u5728\u5404\u9801\u5e95\u90e8\u6587\u5b57\u7dda\u7d22\u4e2d\u5075\u6e2c\u5230\u7a69\u5b9a\u7684\u9801\u78bc\u503c\u3002",
                suggestion="\u82e5 PDF \u6709\u986f\u793a\u9801\u78bc\uff0c\u4f46\u4ecd\u7121\u6cd5\u5075\u6e2c\uff0c\u53ef\u80fd\u662f\u63c3\u63cf\u6216\u7279\u6b8a\u5b57\u578b\u9020\u6210\u3002",
            )
        )

    if not found_catalog:
        issues.append(
            Issue(
                severity="info",
                category="\u7d50\u69cb\u7dda\u7d22",
                title="\u672a\u660e\u78ba\u5075\u6e2c\u5230\u300c\u76ee\u9304\u300d\u5167\u5bb9",
                details="PDF \u6587\u5b57\u64f7\u53d6\u7d50\u679c\u4e2d\u672a\u660e\u78ba\u627e\u5230\u300c\u76ee\u9304\u300d\u95dc\u9375\u5b57\u3002",
            )
        )
    if not found_abstract:
        issues.append(
            Issue(
                severity="info",
                category="\u7d50\u69cb\u7dda\u7d22",
                title="\u672a\u660e\u78ba\u5075\u6e2c\u5230\u6458\u8981\u9801\u7dda\u7d22",
                details="PDF \u6587\u5b57\u64f7\u53d6\u7d50\u679c\u4e2d\u672a\u660e\u78ba\u627e\u5230\u300c\u6458\u8981\u300d\u6216\u300cABSTRACT\u300d\u3002",
            )
        )
    if not found_reference:
        issues.append(
            Issue(
                severity="info",
                category="\u7d50\u69cb\u7dda\u7d22",
                title="\u672a\u660e\u78ba\u5075\u6e2c\u5230\u53c3\u8003\u6587\u737b\u7dda\u7d22",
                details="PDF \u6587\u5b57\u64f7\u53d6\u7d50\u679c\u4e2d\u672a\u660e\u78ba\u627e\u5230\u300c\u53c3\u8003\u6587\u737b\u300d\u3002",
            )
        )
    if not found_chapter:
        issues.append(
            Issue(
                severity="info",
                category="\u7d50\u69cb\u7dda\u7d22",
                title="\u672a\u660e\u78ba\u5075\u6e2c\u5230\u7ae0\u7bc0\u6a19\u984c\u7dda\u7d22",
                details="PDF \u6587\u5b57\u64f7\u53d6\u7d50\u679c\u4e2d\u672a\u660e\u78ba\u627e\u5230\u300c\u7b2cX\u7ae0\u300d\u6a23\u5f0f\u7684\u7ae0\u7bc0\u6a19\u984c\u3002",
            )
        )

    severity_rank = {"error": 0, "warning": 1, "info": 2}
    issues_sorted = sorted(
        issues,
        key=lambda item: (severity_rank.get(item.severity, 9), item.category, item.location or ""),
    )

    page_label_format = "\u672a\u78ba\u5b9a"
    if detected_labels:
        if any(re.fullmatch(r"[ivxlcdmIVXLCDM]+", value) for value in detected_labels):
            page_label_format = "\u7f85\u99ac\u6578\u5b57\u6216\u6df7\u5408"
        elif all(re.fullmatch(r"\d{1,4}", value) for value in detected_labels):
            page_label_format = "\u963f\u62c9\u4f2f\u6578\u5b57"

    return {
        "file_type": "pdf",
        "file_name": path.name,
        "summary": {
            "errors": sum(1 for item in issues_sorted if item.severity == "error"),
            "warnings": sum(1 for item in issues_sorted if item.severity == "warning"),
            "infos": sum(1 for item in issues_sorted if item.severity == "info"),
            "checked_items": [
                "PDF \u9801\u6578",
                "\u9801\u5c3e\u9801\u78bc\u7dda\u7d22",
                "\u53ef\u64f7\u53d6\u6587\u5b57\u5167\u5bb9",
                "\u6458\u8981\u8207 ABSTRACT \u7dda\u7d22",
                "\u76ee\u9304\u7dda\u7d22",
                "\u7ae0\u7bc0\u6a19\u984c\u7dda\u7d22",
                "\u53c3\u8003\u6587\u737b\u7dda\u7d22",
            ],
            "limitations": [
                "PDF \u7121\u6cd5\u50cf DOCX \u4e00\u6a23\u7cbe\u6e96\u53d6\u5f97 Word \u6a23\u5f0f\uff0c\u56e0\u6b64\u7121\u6cd5\u5b8c\u6574\u6aa2\u67e5\u5b57\u578b\uff0c\u7c97\u9ad4\uff0c\u5c0d\u9f4a\u8207\u884c\u8ddd\u7b49\u7d30\u7bc0\u3002",
                "\u82e5 PDF \u70ba\u63c3\u63cf\u6a94\u6216\u5b57\u578b\u5d4c\u5165\u7279\u6b8a\uff0c\u6587\u5b57\u8207\u9801\u78bc\u53ef\u80fd\u7121\u6cd5\u88ab\u6b63\u78ba\u64f7\u53d6\u3002",
                "PDF \u6a21\u5f0f\u8f03\u9069\u5408\u505a\u7d50\u69cb\u8207\u5167\u5bb9\u7dda\u7d22\u6aa2\u67e5\uff0c\u82e5\u8981\u505a\u683c\u5f0f\u7d1a\u5225\u7684\u7cbe\u6e96\u6bd4\u5c0d\uff0c\u5efa\u8b70\u4e0a\u50b3 DOCX\u3002",
            ],
        },
        "coverage": {
            "paragraph_count": total_lines,
            "section_count": len(reader.pages),
            "page_number": {"present": bool(detected_labels), "format": page_label_format},
            "watermark": False,
            "protection": {"enabled": False, "mode": "\u4e0d\u9069\u7528"},
            "page_count": len(reader.pages),
        },
        "section_results": page_results,
        "paragraph_results": [],
        "issues": [asdict(item) for item in issues_sorted],
    }
